import re
from docx import Document
import io


def _ensure_list(value):
    if isinstance(value, list):
        return value
    if value in (None, ""):
        return []
    if isinstance(value, dict):
        return [value]
    return [value]


def _normalize_entry(entry, key_map=None):
    if isinstance(entry, dict):
        if key_map:
            for alias, target in key_map.items():
                if alias in entry and target not in entry:
                    entry[target] = entry[alias]
        return entry
    if isinstance(entry, str):
        return {"description": entry}
    return {}


def generate_docx_bytes(cv_json):
    try:
        doc = Document()
        personal = cv_json.get("personal_info", {})
        doc.add_heading(personal.get("name", ""), level=1)

        # Contact information with social links
        contact_parts = []
        if personal.get("github"):
            contact_parts.append(f"GitHub: {personal.get('github')}")
        if personal.get("linkedin"):
            contact_parts.append(f"LinkedIn: {personal.get('linkedin')}")
        if personal.get("email"):
            contact_parts.append(f"Email: {personal.get('email')}")
        if personal.get("phone"):
            contact_parts.append(f"Phone: {personal.get('phone')}")
        if personal.get("address"):
            contact_parts.append(f"Location: {personal.get('address')}")
        if personal.get("portfolio"):
            contact_parts.append(f"Portfolio: {personal.get('portfolio')}")
        
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

        # Summary
        summary_text = cv_json.get("summary")
        if summary_text:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(summary_text)

        # Current Role and Responsibilities (first experience)
        experience = _ensure_list(cv_json.get("experience", []))
        experience = [_normalize_entry(exp) for exp in experience]
        
        if experience:
            doc.add_heading("Current Role and Responsibilities", level=2)
            exp = experience[0]
            if exp:
                role = exp.get("role", "") or exp.get("description", "")
                company = exp.get("company", "")
                start_date = exp.get("start_date", "")
                end_date = exp.get("end_date", "") or "Present"
                
                # Role and company header
                p = doc.add_paragraph()
                p.add_run(role).bold = True
                p.add_run(f" at {company}")
                
                # Dates
                doc.add_paragraph(f"{start_date} - {end_date}", style='Intense Quote')
                
                # Description with bullet points
                description = exp.get("description", "")
                if description and description != role:
                    # Check if description has bullet points or numbered list
                    if '\n' in description or '•' in description or '-' in description[:20] or re.match(r'^\d+[.)-]', description.strip()):
                        lines = description.split('\n')
                        for line in lines:
                            cleaned = line.strip()
                            # Remove numbered prefixes like "1.", "2.", "1)", etc.
                            if re.match(r'^\d+[.)-]\s*', cleaned):
                                cleaned = re.sub(r'^\d+[.)-]\s*', '', cleaned)
                            cleaned = cleaned.lstrip('•-*').strip()
                            if cleaned:
                                doc.add_paragraph(cleaned, style='List Bullet')
                    else:
                        doc.add_paragraph(description)

        # Career Progression (remaining experiences)
        if len(experience) > 1:
            doc.add_heading("Career Progression", level=2)
            for exp in experience[1:]:
                if not exp:
                    continue
                role = exp.get("role", "") or exp.get("description", "")
                company = exp.get("company", "")
                start_date = exp.get("start_date", "")
                end_date = exp.get("end_date", "") or "Present"
                
                # Role and company header
                p = doc.add_paragraph()
                p.add_run(role).bold = True
                p.add_run(f" at {company}")
                
                # Dates
                doc.add_paragraph(f"{start_date} - {end_date}", style='Intense Quote')
                
                # Description with bullet points
                description = exp.get("description", "")
                if description and description != role:
                    # Check if description has bullet points or numbered list
                    if '\n' in description or '•' in description or '-' in description[:20] or re.match(r'^\d+[.)-]', description.strip()):
                        lines = description.split('\n')
                        for line in lines:
                            cleaned = line.strip()
                            # Remove numbered prefixes like "1.", "2.", "1)", etc.
                            if re.match(r'^\d+[.)-]\s*', cleaned):
                                cleaned = re.sub(r'^\d+[.)-]\s*', '', cleaned)
                            cleaned = cleaned.lstrip('•-*').strip()
                            if cleaned:
                                doc.add_paragraph(cleaned, style='List Bullet')
                    else:
                        doc.add_paragraph(description)
                
                doc.add_paragraph()  # Space between entries

        # Projects
        doc.add_heading("Projects", level=2)
        projects = _ensure_list(cv_json.get("projects", []))
        projects = [_normalize_entry(proj) for proj in projects]
        for proj in projects:
            if not proj:
                continue
            name = proj.get("project_name", "") or proj.get("title", "") or proj.get("description", "")
            description = proj.get("description", "")
            technologies = proj.get("technologies", "")
            
            if name:
                p = doc.add_paragraph()
                p.add_run(name).bold = True
                if technologies:
                    p.add_run(f" - Technologies: {technologies}")
            
            if description and description != name:
                if '\n' in description or '•' in description:
                    lines = description.split('\n')
                    for line in lines:
                        cleaned = line.strip().lstrip('•-*').strip()
                        if cleaned:
                            doc.add_paragraph(cleaned, style='List Bullet')
                else:
                    doc.add_paragraph(description)
            
            doc.add_paragraph()  # Space between projects

        # Education
        doc.add_heading("Education", level=2)
        education = _ensure_list(cv_json.get("education", []))
        education = [
            _normalize_entry(ed, {"college": "institute"}) for ed in education
        ]
        for ed in education:
            if not ed:
                continue
            degree = ed.get("degree", "") or ed.get("description", "")
            institute = ed.get("institute", "")
            start_year = ed.get("start_year", "")
            end_year = ed.get("end_year", "")
            gpa = ed.get("gpa", "") or ed.get("score", "")
            
            p = doc.add_paragraph()
            p.add_run(degree).bold = True
            p.add_run(f" - {institute}")
            if start_year or end_year:
                p.add_run(f" ({start_year} - {end_year})")
            if gpa:
                doc.add_paragraph(f"GPA: {gpa}", style='Intense Quote')

        # Skills
        doc.add_heading("Skills", level=2)
        skills = cv_json.get("skills", [])
        if isinstance(skills, dict):
            # Categorized skills
            for category, skills_list in skills.items():
                # Skip empty General category
                if category == 'General' and (not skills_list or len(skills_list) == 0):
                    continue
                if skills_list:
                    skills_str = skills_list if isinstance(skills_list, str) else ", ".join(str(s) for s in skills_list)
                    p = doc.add_paragraph()
                    p.add_run(f"{category}: ").bold = True
                    p.add_run(skills_str)
        else:
            # Flat list of skills
            skills = _ensure_list(skills)
            flattened_skills = []
            for skill in skills:
                if isinstance(skill, str):
                    flattened_skills.append(skill)
                elif isinstance(skill, dict):
                    flattened_skills.extend(
                        v for v in skill.values() if isinstance(v, str)
                    )
            if flattened_skills:
                doc.add_paragraph(", ".join(flattened_skills))

        doc.add_heading("Certifications", level=2)
        certs = _ensure_list(cv_json.get("certifications", []))
        certs = [_normalize_entry(cert) for cert in certs]
        for cert in certs:
            if not cert:
                continue
            name = cert.get("name", "") or cert.get("description", "")
            issuer = cert.get("issuer", "")
            year = cert.get("year", "")
            line = name
            if issuer or year:
                suffix = " - ".join(filter(None, [issuer, year]))
                line = f"{name} - {suffix}"
            doc.add_paragraph(line)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
    except Exception as e:
        raise Exception(f"Failed to generate DOCX: {str(e)}")
